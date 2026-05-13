"""
Document parser — extracts plain text from uploaded files.

Supports: PDF, DOCX, XLSX, and plain text files.
"""

import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger("nexus.parser")


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    text: str
    filename: str
    page_count: int = 1
    metadata: dict = field(default_factory=dict)


def parse_pdf(content: bytes, filename: str) -> ParsedDocument:
    """Extract text from a PDF file."""
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages_text: list[str] = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(text.strip())

    return ParsedDocument(
        text="\n\n".join(pages_text),
        filename=filename,
        page_count=len(reader.pages),
        metadata={"format": "pdf"},
    )


def parse_docx(content: bytes, filename: str) -> ParsedDocument:
    """Extract text from a DOCX file (paragraphs + tables)."""
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return ParsedDocument(
        text="\n\n".join(parts),
        filename=filename,
        page_count=1,
        metadata={"format": "docx"},
    )


def parse_xlsx(content: bytes, filename: str) -> ParsedDocument:
    """Extract text from an XLSX file (sheet by sheet, row by row)."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}")

        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell) if cell is not None else "" for cell in row]
            row_text = " | ".join(v for v in row_values if v)
            if row_text:
                parts.append(row_text)

    wb.close()
    return ParsedDocument(
        text="\n\n".join(parts),
        filename=filename,
        page_count=len(wb.sheetnames),
        metadata={"format": "xlsx"},
    )


def parse_text(content: bytes, filename: str) -> ParsedDocument:
    """Handle plain text, CSV, and other text-based files."""
    text = content.decode("utf-8", errors="replace")
    return ParsedDocument(
        text=text,
        filename=filename,
        page_count=1,
        metadata={"format": "text"},
    )


# ── Dispatcher ─────────────────────────────────────────────────────────────


# Map of supported MIME types / extensions to parsers
_PARSERS = {
    "application/pdf": parse_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": parse_docx,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": parse_xlsx,
    "application/vnd.ms-excel": parse_xlsx,
}

# Extension fallback
_EXT_PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".xls": parse_xlsx,
    ".txt": parse_text,
    ".csv": parse_text,
    ".md": parse_text,
    ".json": parse_text,
    ".xml": parse_text,
    ".html": parse_text,
    ".htm": parse_text,
}


def parse_document(content: bytes, filename: str, content_type: str = "") -> ParsedDocument:
    """
    Parse a document and return its text content.

    Tries MIME type first, then falls back to file extension.
    """
    # Try MIME type
    parser = _PARSERS.get(content_type)

    # Fallback to extension
    if parser is None:
        import os
        ext = os.path.splitext(filename)[1].lower()
        parser = _EXT_PARSERS.get(ext)

    if parser is None:
        logger.warning(f"Unsupported file type: {filename} ({content_type})")
        # Best-effort: try as plain text
        return parse_text(content, filename)

    try:
        return parser(content, filename)
    except Exception as e:
        logger.error(f"Failed to parse {filename}: {e}")
        raise ValueError(f"Could not parse file '{filename}': {e}") from e
